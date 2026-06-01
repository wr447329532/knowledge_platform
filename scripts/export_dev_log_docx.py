#!/usr/bin/env python3
"""从 README.md 导出「功能与更新记录」为 Word 文档。"""
from __future__ import annotations

import re
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
README = ROOT / "README.md"
OUTPUT = ROOT / "docs" / "文件共享和知识管理平台_开发日志.docx"


def parse_dev_log_sections(text: str) -> list[tuple[str, list[tuple[str, str]]]]:
    """返回 [(章节标题, [(日期, 说明), ...]), ...]"""
    start = text.find("功能与更新记录")
    end = text.find("### 待完善功能（开发计划）")
    if start < 0:
        raise SystemExit("未找到「功能与更新记录」章节")
    block = text[start:end] if end > start else text[start:]

    sections: list[tuple[str, list[tuple[str, str]]]] = []
    current_title: str | None = None
    current_rows: list[tuple[str, str]] = []

    for line in block.splitlines():
        m = re.match(r"^### (.+)$", line.strip())
        if m:
            if current_title and current_rows:
                sections.append((current_title, current_rows))
            current_title = m.group(1).strip()
            current_rows = []
            continue
        row = re.match(r"^\|\s*(\d{4}-\d{2}-\d{2})\s*\|\s*(.+?)\s*\|$", line.strip())
        if row and current_title:
            desc = row.group(2).strip()
            desc = re.sub(r"\*\*(.+?)\*\*", r"\1", desc)
            desc = desc.replace("`", "")
            current_rows.append((row.group(1), desc))

    if current_title and current_rows:
        sections.append((current_title, current_rows))
    return sections


def strip_md_inline(s: str) -> str:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    return s.replace("`", "")


def build_docx(sections: list[tuple[str, list[tuple[str, str]]]], out_path: Path) -> None:
    try:
        from docx import Document
    except ImportError as e:
        raise SystemExit(
            "缺少依赖 python-docx。请在当前 Python 环境中执行:\n"
            "  pip install python-docx"
        ) from e
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "PingFang SC"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    title = doc.add_heading("文件共享和知识管理平台 — 开发日志", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("导出自项目 README.md · 功能与更新记录")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    for section_title, rows in sections:
        doc.add_heading(section_title, level=1)
        table = doc.add_table(rows=1 + len(rows), cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "时间"
        hdr[1].text = "功能说明"
        for i, (dt, desc) in enumerate(rows, start=1):
            table.rows[i].cells[0].text = dt
            table.rows[i].cells[1].text = desc
        doc.add_paragraph()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    text = README.read_text(encoding="utf-8")
    sections = parse_dev_log_sections(text)
    if not sections:
        raise SystemExit("未解析到任何开发日志条目")
    build_docx(sections, OUTPUT)
    total = sum(len(r) for _, r in sections)
    print(f"已生成: {OUTPUT}")
    print(f"共 {len(sections)} 个章节, {total} 条记录")


if __name__ == "__main__":
    main()
